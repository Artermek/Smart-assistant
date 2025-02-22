from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
import time
import re
from langchain.document_loaders import AsyncHtmlLoader
from langchain.document_transformers import Html2TextTransformer
from openai import OpenAI # работа с API OpenAI
import re                 # работа с регулярными выражениями
import os
from dotenv import load_dotenv
from docx import Document
from PIL import Image, ImageDraw, ImageFont
import requests
# получим переменные окружения из .env
load_dotenv()

# API-key
openai.api_key = os.environ.get("OPENAI_API_KEY")



def search_google_news(query):
    chrome_options = Options()
    chrome_options.add_argument("--headless")  # Без графического интерфейса
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/90.0.4430.212 Safari/537.36")

    service = Service('/usr/bin/chromedriver')
    driver = webdriver.Chrome(service=service, options=chrome_options)

    try:
        # Открываем страницу поиска Google News
        url = f"https://news.google.com/search?q={query}"
        driver.get(url)

        # Ожидаем загрузку статей
        WebDriverWait(driver, 15).until(
            EC.presence_of_element_located((By.XPATH, '//article'))
        )

        # Получаем HTML-код страницы
        page_source = driver.page_source
        soup = BeautifulSoup(page_source, 'html.parser')
        articles = soup.find_all('article')[:4]  # Берем 5 новостей

        news = []
        for article in articles:
            title_element = article.find('a')
            if title_element and 'href' in title_element.attrs:
                google_news_link = f"https://news.google.com{title_element['href'].lstrip('.')}"

                # Переход на страницу новости
                driver.get(google_news_link)

                # Ждем загрузки реального сайта
                time.sleep(7)  # Можно увеличить до 10 сек, если сайт грузится долго

                # Получаем реальный URL после редиректа
                real_link = driver.current_url

                title = title_element.text.strip()
                news.append({"title": title, "link": real_link})

        return news

    except Exception as e:
        print(f"Ошибка: {str(e)}")
        return []

    finally:
        driver.quit()

def clean_text(text):

  # удаление заголовков и подзаголовков
  text = re.sub(r"\* .+\n", "", text)

  # удаление разделительных линий и специальных символов
  text = re.sub(r"__+", "", text)

  # удаление ссылок и инструкций JavaScript
  text = re.sub(r"Нажимая кнопку .+\n", "", text)
  text = re.sub(r"Пожалуйста, включите JavaScript .+\n", "", text)

  # удаление строк типа "Оценка квартиры для [Название Банка] Подробнее"
  text = re.sub(r"Оценка квартиры для .+ Подробнее \n", "", text)

  # удаление строк, содержащих "Заказать звонок"
  text = re.sub(r".*Заказать звонок.*\n", "", text)

  # очистка текста от лишних пробелов и переводов строк
  text = re.sub(r"\n\s*\n", "\n", text)

  # возврат результата
  return text



system = """Ты проффесионально пишешь тексты для блога. Твоя задача на основе предоставленного текста создать краткий,
но содержательный текст, подходящий для блога. Текст должен быть написан простым и понятным языком, рассчитанным на аудиторию мам.
"""

user = """ Напиши краткий, но содержательный текст для блога на основе новостей. Текст должен быть рассчитан на аудиторию мам и говорить о важности раннего обучения детей.
Текст должен быть написан простым и понятным языком, рассчитанным на аудиторию мам.

 """


def generate_blog_text(system, user, news):
    client = OpenAI()

    messages = [
        {"role": "system", "content": system},
        {"role": "user", "content": """
Как посмотреть, где находится устройство ребенка.Для того чтобы точно определить местоположение ребенка с помощью телефона, необходимо использовать специальные приложения. Эти приложения позволяют отслеживать передвижения с помощью GPS и других технологий. Больше всего вам подойдут сервисы, которые работают как на Android, так и на iOS. Кроме установки приложения, важно убедиться, что на устройстве ребенка включены основные функции геолокации. Для этого можно выполнить следующие действия: Проверьте настройки геолокации в меню устройства. Убедитесь, что у приложения есть доступ к определению местоположения. Настройте параметры уведомлений, чтобы получать сигналы о перемещениях ребенка. Обратите внимание также на то, что большинство приложений позволяет делиться своими координатами с другими пользователями.
        """},
        {"role": "assistant", "content": """Фотореалистичное изображение маленькой девочки (примерно 6–7 лет) со светло-коричневыми волосами, собранными в два аккуратных пучка.  Она сидит на полу в уютной, светлой комнате рядом с большой картонной коробкой. На ней надета синяя клетчатая рубашка поверх белой футболки. Девочка держит в руках смартфон и плюшевого медвежонка, при этом выглядит увлечённой и сосредоточенной. Естественное освещение, мягкие тени, тёплая домашняя  атмосфера. Высокая детализация, фотореалистичный стиль. на изображение не должно быть текста!
        """},
        {"role": "user", "content": f"{user} \nНовости:{news}"}
    ]

    completion = client.chat.completions.create(
        model="gpt-4o-mini",
        messages= messages,
        temperature=0.5
    )

    return completion.choices[0].message.content



system_img = """Ты профессионально пишешь промты для модели Dalle-3 . Тебе нужно из предоставленного текста ,
для блога  выделить основную мысль или лозунг, который будет использоваться для создания изображения например,
баннера или инфографики, на твой выбор. Текст должен быть кратким и запоминающимся
"""
user_img = """ Твоя задача написать промт для создания изображения для Dalle-3. Тебе будет дан текст из которого
тебе нужно выделить основную мысль или лозунг который будет использоваться для создания изображения например,
баннера или инфографики, на твой выбор. Текст должен быть кратким и запоминающимся. В промт обязательно добавь, что картинка должна быть на русском языке!.
Не пиши что это промт для Dalle-3. Напиши только промт, не добавляй ничего лишнего
"""


def generate_img_text(system, user, final_text):

    client = OpenAI()

    messages = [
        {"role": "system", "content": system},
        {"role": "user", "content": f"{user} \nТекст:{final_text}"}
    ]

    completion = client.chat.completions.create(
        model="gpt-4o-mini",
        messages= messages,
        temperature=0.5
    )

    return completion.choices[0].message.content



def generate_image_with_dalle(prompt, output_path="output_image.png"):
    # Инициализация клиента OpenAI
    client = OpenAI()

    # Генерация изображения
    response = client.images.generate(
        model="dall-e-3",  # Используем DALL-E 3
        prompt=prompt,
        size="1024x1024",  # Размер изображения
        quality="standard",  # Качество изображения
        n=1  # Количество изображений
    )

    # Сохранение изображения
    image_url = response.data[0].url
    print(f"Изображение сгенерировано: {image_url}")

    # Скачивание изображения (опционально)
    import requests
    image_data = requests.get(image_url).content
    with open(output_path, "wb") as f:
        f.write(image_data)
    print(f"Изображение сохранено в {output_path}")


def save_blog_text(blog_text):
    doc = Document()
    doc.add_paragraph(blog_text)
    doc.save("blog_text.docx")

def save_img_text(img_text):
    doc = Document()
    doc.add_paragraph(img_text)
    doc.save("img_text.docx")




def main():
  # Пример использования
  query = "early teaching of children arithmetic and speed reading"
  news = search_google_news(query)

  #print("Найденные новости:")
  site_news = []
  for i, item in enumerate(news, start=1):
      site_news.append(item['link'])
      #print(f"{item['link']}\n")

  print(site_news[0])

  # создание объекта загрузчика
  loader = AsyncHtmlLoader(site_news[0])

  # загружаем текст
  docs = loader.load()


  # создаем объект для трасформации в текст
  html2text = Html2TextTransformer()

  # трансформируем в текст
  docs_transformed = html2text.transform_documents(docs)

  for doc in docs_transformed:
    doc.page_content = clean_text(doc.page_content)

  # посмотреть очищенный текст
  #for doc in docs_transformed:
    #print(doc.page_content)

  final_text = generate_blog_text(system, user, doc.page_content)
  #print (final_text)

  promt_img = generate_img_text (system_img, user_img, final_text)
  #print (promt_img)

  generate_image_with_dalle(promt_img)

  save_blog_text(final_text)
  save_img_text (promt_img)

  print("Ассистент завершил работу. Файлы сохранены.")

if __name__ == "__main__":
    main()





